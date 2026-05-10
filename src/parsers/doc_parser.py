from dataclasses import dataclass, field
from pathlib import Path
from loguru import logger


@dataclass
class DocChunk:
    text: str
    source: str
    page_index: int    # 0-based (PDF: 페이지, DOCX: 단락 그룹)
    doc_type: str      # "rfp" | "proposal"
    metadata: dict = field(default_factory=dict)

    @property
    def slide_index(self) -> int:
        return self.page_index


# ── PDF ──────────────────────────────────────────────────────────────────────

def _parse_pdf(path: Path, doc_type: str) -> list[DocChunk]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber 미설치: pip install pdfplumber")

    chunks: list[DocChunk] = []
    with pdfplumber.open(str(path)) as pdf:
        for idx, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            text = text.strip()
            if not text:
                continue

            # 테이블도 추출해서 합치기
            tables = page.extract_tables() or []
            for table in tables:
                for row in table:
                    row_text = " | ".join(str(cell or "").strip() for cell in row if cell)
                    if row_text.strip():
                        text += "\n" + row_text

            chunks.append(DocChunk(
                text=text,
                source=str(path),
                page_index=idx,
                doc_type=doc_type,
                metadata={"file_name": path.name, "total_pages": len(pdf.pages)},
            ))

    logger.info(f"PDF 파싱 완료: {path.name} — {len(chunks)} 페이지")
    return chunks


# ── DOCX ─────────────────────────────────────────────────────────────────────

def _parse_docx(path: Path, doc_type: str) -> list[DocChunk]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx 미설치: pip install python-docx")

    doc = Document(str(path))
    chunks: list[DocChunk] = []
    current_lines: list[str] = []
    group_idx = 0

    def flush(lines: list[str], idx: int):
        text = "\n".join(lines).strip()
        if text:
            chunks.append(DocChunk(
                text=text,
                source=str(path),
                page_index=idx,
                doc_type=doc_type,
                metadata={"file_name": path.name},
            ))

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 헤딩이면 새 그룹 시작
        is_heading = para.style.name.startswith("Heading")
        if is_heading and current_lines:
            flush(current_lines, group_idx)
            current_lines = []
            group_idx += 1

        current_lines.append(text)

    # 테이블 텍스트도 추출
    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
        if table_lines:
            flush(table_lines, group_idx)
            group_idx += 1

    flush(current_lines, group_idx)

    logger.info(f"DOCX 파싱 완료: {path.name} — {len(chunks)} 그룹")
    return chunks


# ── 공통 진입점 ────────────────────────────────────────────────────────────────

def parse_document(file_path: str | Path, doc_type: str = "rfp") -> list[DocChunk]:
    """PDF 또는 DOCX 파일 파싱. doc_type: 'rfp' | 'proposal'"""
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"파일 없음: {path}")

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _parse_pdf(path, doc_type)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(path, doc_type)
    else:
        raise ValueError(f"지원하지 않는 형식: {suffix} (pdf, docx만 지원)")


def parse_rfp_dir(dir_path: str | Path, doc_type: str = "rfp") -> list[DocChunk]:
    """디렉토리 내 모든 PDF/DOCX 파일 일괄 파싱."""
    dir_path = Path(dir_path)
    files = (
        list(dir_path.glob("**/*.pdf"))
        + list(dir_path.glob("**/*.PDF"))
        + list(dir_path.glob("**/*.docx"))
        + list(dir_path.glob("**/*.DOCX"))
    )
    if not files:
        logger.warning(f"PDF/DOCX 파일 없음: {dir_path}")
        return []

    all_chunks: list[DocChunk] = []
    for f in files:
        try:
            all_chunks.extend(parse_document(f, doc_type))
        except Exception as e:
            logger.error(f"파싱 실패 {f.name}: {e}")

    logger.info(f"총 {len(all_chunks)}개 문서 청크 추출 ({len(files)}개 파일)")
    return all_chunks
